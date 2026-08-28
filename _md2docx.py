import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = r"c:\Users\wuzhixian\Desktop\Excel-Agent-And-Merge\docs\答辩PPT内容文档-45min.md"
DST = r"c:\Users\wuzhixian\Desktop\Excel-Agent-And-Merge\docs\答辩PPT内容文档-45min.docx"

NAVY = RGBColor(0x1F, 0x4E, 0x79)
ORANGE = RGBColor(0xED, 0x7D, 0x31)
GRAY = RGBColor(0x59, 0x59, 0x59)
BLACK = RGBColor(0x00, 0x00, 0x00)

with open(SRC, encoding="utf-8") as f:
    lines = f.read().splitlines()

doc = Document()

# base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def set_cn_font(run, name="Microsoft YaHei"):
    run.font.name = name
    r = run._element.rPr.rFonts
    r.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", name)

def add_code_block(code_lines):
    for cl in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(cl if cl else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        run.font.color.rgb = GRAY
        # light gray shading
        from docx.oxml.ns import qn
        shd = p._p.get_or_add_pPr()
        el = shd.find(qn('w:shd'))
        if el is None:
            el = shd.makeelement(qn('w:shd'), {})
            shd.append(el)
        el.set(qn('w:fill'), 'F2F2F2')

def add_table(rows):
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncol)
    tbl.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncol):
            cell = tbl.cell(i, j)
            txt = row[j] if j < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            set_cn_font(run)
            run.font.size = Pt(9.5)
            if i == 0:
                run.font.bold = True
                run.font.color.rgb = NAVY
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

i = 0
n = len(lines)
in_code = False
code_buf = []
in_table = False
table_buf = []

def flush_code():
    global in_code, code_buf
    if code_buf:
        add_code_block(code_buf)
        code_buf = []
    in_code = False

def flush_table():
    global in_table, table_buf
    if table_buf:
        add_table(table_buf)
        table_buf = []
    in_table = False

def parse_row(line):
    s = line.strip()
    if s.startswith("|"):
        s = s.strip("|")
        cells = [c.strip() for c in s.split("|")]
        return cells
    return None

while i < n:
    line = lines[i]
    stripped = line.strip()

    # code fence
    if stripped.startswith("```"):
        if in_code:
            flush_code()
        else:
            flush_table()
            in_code = True
            code_buf = []
        i += 1
        continue

    if in_code:
        code_buf.append(line)
        i += 1
        continue

    # table
    if stripped.startswith("|"):
        row = parse_row(line)
        # skip separator rows
        if row and all(re.fullmatch(r":?-{2,}:?", c) for c in row):
            i += 1
            continue
        if row:
            if not in_table:
                flush_code()
                in_table = True
                table_buf = []
            table_buf.append(row)
        i += 1
        continue
    else:
        flush_table()

    # headings
    if stripped.startswith("# "):
        flush_code()
        p = doc.add_heading(stripped[2:], level=0)
        for r in p.runs:
            set_cn_font(r)
            r.font.color.rgb = NAVY
    elif stripped.startswith("## "):
        flush_code()
        p = doc.add_heading(stripped[3:], level=1)
        for r in p.runs:
            set_cn_font(r)
            r.font.color.rgb = NAVY
    elif stripped.startswith("### "):
        flush_code()
        p = doc.add_heading(stripped[4:], level=2)
        for r in p.runs:
            set_cn_font(r)
            r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    elif stripped.startswith("> "):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        run = p.add_run(stripped[2:])
        set_cn_font(run)
        run.italic = True
        run.font.color.rgb = GRAY
    elif stripped.startswith("---"):
        pass  # skip hr
    elif stripped == "":
        pass
    else:
        # normal paragraph, handle inline bold
        p = doc.add_paragraph()
        text = stripped
        # bullet list
        if re.match(r"^[-*] ", text):
            text = text[2:]
            p.style = doc.styles["List Bullet"]
        elif re.match(r"^\d+\. ", text):
            text = re.sub(r"^\d+\. ", "", text, count=1)
            p.style = doc.styles["List Number"]
        # bold segments
        pos = 0
        for m in re.finditer(r"\*\*(.+?)\*\*", text):
            if m.start() > pos:
                r1 = p.add_run(text[pos:m.start()])
                set_cn_font(r1)
            r2 = p.add_run(m.group(1))
            set_cn_font(r2)
            r2.bold = True
            pos = m.end()
        if pos < len(text):
            r3 = p.add_run(text[pos:])
            set_cn_font(r3)
    i += 1

flush_code()
flush_table()

doc.save(DST)
print("saved:", DST)
