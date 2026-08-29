import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = r"c:\Users\wuzhixian\Desktop\Excel-Agent-And-Merge\docs\Merge部分详细萃取.md"
OUT = r"c:\Users\wuzhixian\Desktop\Merge-PPT每页组织方案.docx"

START_MARK = "# 第 9 章：PPT 每页组织方案"
END_MARK = "# 附录：术语速查表"


def set_font(run, name="微软雅黑", size=10.5, bold=False, color=None, mono=False):
    fname = "Consolas" if mono else name
    run.font.name = fname
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    if mono:
        r = run._element
        r.rPr.rFonts.set(qn("w:ascii"), fname)
        r.rPr.rFonts.set(qn("w:hAnsi"), fname)
    else:
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    sizes = {1: 16, 2: 13, 3: 11.5}
    run = p.add_run(text)
    set_font(run, size=sizes.get(level, 11), bold=True,
             color=RGBColor(0x1F, 0x3B, 0x63) if level <= 2 else RGBColor(0x33, 0x33, 0x33))


def add_inline_text(p, text):
    # 处理 `code` 与 **bold** 内联
    pos = 0
    pattern = re.compile(r'(\*\*.*?\*\*|`[^`]*`)')
    for m in pattern.finditer(text):
        if m.start() > pos:
            set_font(p.add_run(text[pos:m.start()]))
        seg = m.group(0)
        if seg.startswith('**') and seg.endswith('**'):
            set_font(p.add_run(seg[2:-2]), bold=True)
        else:
            set_font(p.add_run(seg[1:-1]), mono=True)
        pos = m.end()
    if pos < len(text):
        set_font(p.add_run(text[pos:]))


def add_para(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Pt(18)
    add_inline_text(p, text)
    return p


def add_code_block(doc, code):
    for line in code.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Pt(12)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line if line else " ")
        set_font(run, mono=True, size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, header, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.paragraphs[0].text = ""
        r = cell.paragraphs[0].add_run(h)
        set_font(r, bold=True, size=9.5)
    for i, row in enumerate(rows):
        for j, cellval in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.paragraphs[0].text = ""
            add_inline_text(cell.paragraphs[0], cellval)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def render_table(doc, lines):
    rows = []
    for ln in lines:
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    header = rows[0]
    body = rows[2:] if len(rows) > 1 and all(
        re.fullmatch(r":?-{3,}:?", c.strip()) for c in rows[1]) else rows[1:]
    add_table(doc, header, body)


def main():
    with open(SRC, encoding="utf-8") as f:
        text = f.read()
    start = text.index(START_MARK)
    end = text.index(END_MARK, start)
    body = text[start:end].rstrip()

    doc = Document()
    lines = body.split("\n")
    i = 0
    in_code = False
    code_lines = []
    in_table = False
    table_lines = []

    def flush_code():
        nonlocal code_lines, in_code
        if code_lines:
            add_code_block(doc, "\n".join(code_lines))
            code_lines = []
            in_code = False

    def flush_table():
        nonlocal table_lines, in_table
        if table_lines:
            render_table(doc, table_lines)
            table_lines = []
            in_table = False

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                flush_code()
            else:
                flush_table()
                in_code = True
                code_lines = []
            continue
        if in_code:
            code_lines.append(line)
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(stripped)
            continue
        else:
            flush_table()

        if stripped.startswith("## "):
            flush_code()
            add_heading(doc, stripped[3:].strip(), 2)
        elif stripped.startswith("# "):
            flush_code()
            add_heading(doc, stripped[2:].strip(), 1)
        elif stripped.startswith("### "):
            flush_code()
            add_heading(doc, stripped[4:].strip(), 3)
        elif stripped == "---":
            continue
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Pt(12)
            r = p.add_run(stripped[2:].strip())
            set_font(r, size=9.5, color=RGBColor(0x66, 0x66, 0x66))
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            add_inline_text(p, stripped[2:].strip())
        elif stripped == "":
            continue
        else:
            add_para(doc, stripped)

    flush_code()
    flush_table()

    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    main()
